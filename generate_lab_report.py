from docx import Document

doc = Document()

doc.add_heading('JOKSHU Voting System - Lab Report', level=1)

doc.add_heading('1. Introduction', level=2)
doc.add_paragraph(
    'JOKSHU is a mobile voting application designed for Jagannath University student union elections. '
    'Built with Expo, React Native and Firebase, the app supports secure student login, candidate browsing, '
    'vote casting, and administrative election management. The system is optimized for an APK-based '
    'deployment, allowing users to install the app directly on Android devices without requiring QR scan or local network access.'
)

doc.add_heading('2. Motivation', level=2)
doc.add_paragraph(
    'The primary motivation for JOKSHU is to modernize campus election management while preserving security and usability. '
    'The application aims to remove the friction of traditional ballot systems, provide a secure voting flow, and enable '
    'administrators to review election progress in real time. Special emphasis is placed on securing student identity verification and protecting voter data.'
)

doc.add_heading('3. Requirements', level=2)
doc.add_paragraph('Functional requirements:')
doc.add_paragraph(
    '- Student authentication with Student ID and password\n'
    '- Candidate browsing and position-based voting\n'
    '- Vote tracking and result display\n'
    '- Admin controls for election state and OTP management\n'
    '- ID card verification and optional face/OTP verification for entry'
)

doc.add_paragraph('Non-functional requirements:')
doc.add_paragraph(
    '- Mobile-first responsive user interface\n'
    '- Secure password hashing and OTP validation\n'
    '- Offline-capable APK installation for direct deployment\n'
    '- Stable performance on Expo SDK 54 and Android devices'
)

doc.add_heading('4. System Design and Development', level=2)
doc.add_paragraph(
    'The JOKSHU app is structured as a React Native Expo project with a context-based state management layer. '
    'The architecture includes separate modules for authentication, voting logic, navigation, and device features like camera scanners.'
)

doc.add_paragraph('Key components:')
doc.add_paragraph(
    '- AuthContext: manages login, password hashing, and current user storage.\n'
    '- VotingContext: tracks candidates, election state, vote submission, and OTP request flow.\n'
    '- Screens: Login, Home, Voting, Results, Profile, Admin, GPA utilities, and extras.\n'
    '- Components: IDCardScanner, FaceScanner, FingerprintScanner, BMI Calculator, Scientific Calculator, and Snake Game.'
)

doc.add_paragraph(
    'The app uses Expo and Firebase for persistent storage and cloud functions. Local storage is handled through AsyncStorage and a FirebaseStorage abstraction. '
    'The OTP verification process employs a Firebase Cloud Function that generates a random 6-digit code and sends it using Twilio SMS. '
    'Election data is stored and updated locally for votes, while administrative updates are persisted in Firebase storage.'
)

doc.add_heading('4.1 Security', level=3)
doc.add_paragraph(
    'Security is a central feature of the JOKSHU system. The app uses hashed passwords, secure storage, and multi-factor verification mechanisms. '
    'Passwords are hashed using SHA-256 before storage, and custom password entries are re-hashed if needed. '
    'The voting flow requires identity verification using QR-based student ID scanning, optional face recognition, and optional OTP approval from the admin.'
)

doc.add_paragraph(
    'The OTP service is implemented in Firebase Cloud Functions, which validates the request and stores a hashed version of the OTP code. '
    'The actual OTP message sent to the user contains only the minimal necessary details: the one-time code and expiration time. '
    'This reduces the risk of sensitive data exposure over SMS. The cloud function also enforces a maximum number of attempts and expiration time to prevent reuse.'
)

doc.add_paragraph(
    'Admin-level access is protected by separate credentials and allows only authorized users to approve or reject OTP requests, reset vote data, and manage election state. '
    'This ensures a layered security model where both the voter and the admin participate in the verification process.'
)

doc.add_heading('5. Output', level=2)
doc.add_paragraph(
    'The final application delivers a polished Android voting app with the following outputs:'
)
doc.add_paragraph(
    '- A deployable APK that can be installed directly on Android devices.\n'
    '- A secure login interface for students and admins.\n'
    '- A verified voting process with ID scanning, OTP approval, and vote confirmation.\n'
    '- Real-time results and admin dashboards for election monitoring.\n'
    '- Additional utility screens for GPA checking, GPA ranking, BMI calculation, and a simple game for user engagement.'
)

doc.add_heading('6. Conclusion', level=2)
doc.add_paragraph(
    'JOKSHU Voting System provides a secure and accessible mobile platform for student elections. '
    'The app combines Expo-based mobile development with Firebase-backed storage and SMS-based OTP verification to create a robust voting experience. '
    'By focusing on direct APK deployment, the project removes dependency on Expo Go and local network access, making it easy to distribute and install on student devices.'
)

doc.add_paragraph(
    'The final system demonstrates how modern mobile app development can improve campus governance while maintaining strong security and administrative control. '
    'The included biometric and OTP safeguards help ensure that only verified students can participate in the vote.'
)

doc.save('JOKSHU_Lab_Report.docx')
print('generated')
