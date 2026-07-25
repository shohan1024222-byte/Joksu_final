from docx import Document

doc = Document()

doc.add_heading('NotebookLM PPT Prompt for JOKSHU Voting App', level=1)

doc.add_paragraph(
    'Use this prompt in NotebookLM to generate a 15-slide PowerPoint presentation for the JOKSHU mobile voting application. ' 
    'The presentation should reflect the app features, design, development approach, output, and security measures. ' 
    'Keep the language professional, clear, and suitable for an academic or project demonstration setting.'
)

doc.add_heading('Prompt', level=2)
doc.add_paragraph(
    'Create a professional 15-slide PowerPoint presentation for the JOKSHU Voting System mobile app. ' 
    'The app is built with Expo, React Native, TypeScript, Firebase, and EAS, and it is intended for Jagannath University student union elections. ' 
    'The slides should cover the app introduction, motivation, functional and non-functional requirements, system architecture, app features, user flow, admin capabilities, build and deployment, output, and security. ' 
    'The security slide should explain password hashing, OTP verification, admin approval workflow, and student ID verification without overstating the technical depth.'
)

doc.add_heading('Slide Structure', level=2)

doc.add_paragraph('1. Title slide')

doc.add_paragraph('2. Project overview and objective')

doc.add_paragraph('3. Motivation and problem statement')

doc.add_paragraph('4. Technology stack and tools (Expo, React Native, Firebase, EAS)')

doc.add_paragraph('5. User authentication and login flow')

doc.add_paragraph('6. Candidate browsing and position-wise voting')

doc.add_paragraph('7. Voting screen and vote confirmation')

doc.add_paragraph('8. Admin dashboard and election controls')

doc.add_paragraph('9. Identity verification features: QR ID scanning, face scanner, fingerprint support')

doc.add_paragraph('10. OTP-based verification process and admin approval')

doc.add_paragraph('11. Additional utilities: GPA tools, BMI calculator, scientific calculator, snake game')

doc.add_paragraph('12. Build and deployment: APK generation with EAS and standalone app installation')

doc.add_paragraph('13. Output and expected user experience')

doc.add_paragraph('14. Security slide: password hashing, secured OTP service, admin control, verification flow')

doc.add_paragraph('15. Conclusion and future improvement suggestions')

doc.add_heading('Additional Instructions for NotebookLM', level=2)

doc.add_paragraph(
    '• Use concise bullet points on each slide.\n'
    '• Include a visual flow description for the voting verification process.\n'
    '• Emphasize that the app is APK-ready for direct installation without QR or local host dependency.\n'
    '• Present the security slide as "Security and Verification" with practical safeguards rather than deep cryptographic detail.\n'
    '• Keep the presentation suitable for a student project defense.'
)

doc.save('NotebookLM_PPT_Prompt.docx')
print('created')
